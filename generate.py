import os

import docx
import csv
from operator import itemgetter

import config


def read_csv(path):
    """
    Чтение csv-файлов.
    :param path: путь до файла
    :return: список словарей: дата, человек, наряд, группа
    """
    with open(path, 'r', encoding='utf-8') as file:
        reader = csv.reader(file, delimiter=';')
        people = []
        klass = path[-7:-4]
        for row in reader:
            man = {'data': int(row[0]), 'man': row[1], 'nar': row[2], 'gr': klass}
            people.append(man)
    return people


def new_list(list_mans):
    """ Преобразование короткой записи в большую запись наряда
    :param list_mans: список словарей людей
    :return: новый список с длинными названиями нарядов
    """
    a = list_mans
    for line in a:
        line['nar'] = config.DICT_FOR_DEZ[line['nar']]
    return a


def list_data_count(list_mans):
    """ Получение списка с датой и числом заступления людей из списка с людьми"""
    data = 0
    list_nar = []
    for man in list_mans:
        if data == man['data']:
            continue
        data = man['data']
        count_data = sum(i.get('data') == data for i in list_mans)
        list_nar.append([data, count_data])
    return list_nar


def generate_heading(document):
    """ Генерим шапку списка нарядов """
    par = document.add_heading("Список", 0)
    par = document.add_heading("Дежурств от 10-х классов лицея №23", 0)
    table = document.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for row in table.rows:
        for i in range(len(row.cells)):
            if i is 0:
                row.cells[i].text = 'Вид дежурства'
                row.cells[i].paragraphs[0].style = 'NumG'
            if i is 1:
                row.cells[i].text = 'Класс'
                row.cells[i].paragraphs[0].style = 'NumG'
            if i is 2:
                row.cells[i].text = 'Ф.И.О.'
                row.cells[i].paragraphs[0].style = 'NumG'
            if i is 3:
                row.cells[i].text = 'Примечание'
                row.cells[i].paragraphs[0].style = 'NumG'


def generate_main(list_nar, document, month_for_generate):
    """
    Функция, непосредственной генерации каждого из дежурств
    :param list_nar: список людей и нарядов из функциия list_data_count
    :param document: документ в который идет генерация
    :param month_for_generate: месяц генерации
    :return: документ document со списком нарядов
    """
    for data in list_nar:
        str_nar = '«' + str(data[0]) + '» ' + month_for_generate + ' 2020 года'
        data_text = document.add_paragraph(str_nar, style='Data')
        table = document.add_table(rows=data[1], cols=4)
        table.style = 'Table Grid'
        mans_in_date = [i for i in mans if i.get('data') == data[0]]
        for row, man in zip(table.rows, mans_in_date):
            for i in range(len(row.cells)):
                if i is 0:
                    row.cells[i].text = man['nar']
                    row.cells[i].paragraphs[0].style = 'A'
                if i is 1:
                    row.cells[i].text = man['gr']
                    row.cells[i].paragraphs[0].style = 'NumG'
                if i is 2:
                    row.cells[i].text = 'уч. ' + man['man']
                    row.cells[i].paragraphs[0].style = 'A'


def generate_list(list_mans, month_for_generate, path):
    """ Основная функция генерации списка дежурств
    :param list_mans: Список словарей людей (получаются из функции get_list_people)
    :param month_for_generate: месяц на который генерируется список
    :param path: путь выходного файла
    """
    people = list_mans
    people.sort(key=itemgetter('data'))  # сортировка по дате дежурства
    list_nar = list_data_count(people)  # список с датами и количеством дежурств

    file = open(config.PATH_TEMPLATE, 'rb')
    doc = docx.Document(file)
    generate_heading(doc)  # генерим заголовки
    generate_main(list_nar, doc, month_for_generate[0])  # генеация остальной части дежурств

    doc.save(path)  # сохранение файла
    pass


def get_list_people(path):
    people = read_csv(path)
    people = new_list(people)
    return people


if __name__ == "__main__":
    print('Генерируем дежурства из файлов с людьми.')
    month = list()
    month.append(input('Введите название месяца в Родительном падеже: '))
    month.append(input('Введите название месяца в Именительном падеже: '))
    mans = []
    for file_path in os.listdir(config.PATH_DIR_DATA):
        mans += get_list_people(config.PATH_DIR_DATA + file_path)

    path_out_file = config.PATH_OUT_FILE_BEGIN + month[1] + config.PATH_OUT_FILE_END
    generate_list(mans, month, path_out_file)  # запуск генерации
    print()
    print('Файл сгенерировался.')
    print("Он доступен по адресу: " + path_out_file)

